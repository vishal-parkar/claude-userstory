from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = "E:/claude_code/UserStories/user-stories"

def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def create_us_doc(filepath, us_id, title, epic, story_type, role, goal, benefit, acceptance_criteria, notes=None):
    doc = Document()

    # Title heading
    title_para = doc.add_heading(f"{us_id} - {title}", level=1)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    # Metadata
    p = doc.add_paragraph()
    p.add_run("Epic: ").bold = True
    p.add_run(epic)

    p = doc.add_paragraph()
    p.add_run("Story ID: ").bold = True
    p.add_run(us_id)

    p = doc.add_paragraph()
    p.add_run("Type: ").bold = True
    p.add_run(story_type)

    p = doc.add_paragraph()
    p.add_run("Priority: ").bold = True
    p.add_run("Medium")

    doc.add_paragraph()
    doc.add_heading("User Story", level=2)

    p = doc.add_paragraph()
    p.add_run(f"As a {role},\n")
    p.add_run(f"I want to {goal},\n")
    p.add_run(f"So that {benefit}.")

    doc.add_paragraph()
    doc.add_heading("Acceptance Criteria", level=2)
    add_bullet_list(doc, acceptance_criteria)

    if notes:
        doc.add_paragraph()
        doc.add_heading("Notes / Clarifications", level=2)
        add_bullet_list(doc, notes)

    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
    print(f"Created: {os.path.basename(filepath)}")

# ─── EPIC 1: Expression Builder Panel ───────────────────────────────────────
E1 = f"{BASE}/Epic-1-Expression-Builder-Panel"

create_us_doc(
    f"{E1}/US-01-Display-Expression-Builder-Main-Screen.docx",
    "US-01", "Display Expression Builder Main Screen",
    "Epic 1 - Expression Builder Panel", "Frontend",
    "user",
    "view the RBP Calculator main screen",
    "I can start building rule-based expressions from a clean interface",
    [
        "The main screen is displayed when the user navigates to the RBP Calculator.",
        "The screen contains an 'Expression' clickable label/button.",
        "The screen contains an 'Add Comparator' clickable label/button.",
        "The screen contains a '+ and/or' button for adding additional conditions.",
        "The screen contains a 'Group' button for creating subgroup conditions.",
        "The layout is clean with no extra dropdowns (e.g., '1Min' dropdown must NOT appear).",
        "The screen shows an 'Add Expression' area/field where expression tokens are displayed.",
    ],
    ["No '1Min' dropdown should be present in the design - explicitly excluded by BA."]
)

create_us_doc(
    f"{E1}/US-02-Open-Expression-Builder-Popup.docx",
    "US-02", "Open Expression Builder Popup on Expression Click",
    "Epic 1 - Expression Builder Panel", "Frontend",
    "user",
    "click on 'Expression' to open the expression builder popup",
    "I can start configuring the expression I want to build",
    [
        "Clicking the 'Expression' label/button opens a popup/modal panel.",
        "The popup contains three tabs or filter options: 'Parameter', 'Math Function', 'All'.",
        "The label 'Indicators' is renamed to 'Parameter' in the popup.",
        "The label 'Math Function' remains unchanged.",
        "The 'All' option label remains unchanged.",
        "The popup displays the correct options based on the selected tab.",
        "The popup is dismissible (e.g., by clicking outside or a close button).",
    ],
    ["Needs Clarification: Is the popup a modal dialog or an inline dropdown panel?"]
)

# ─── EPIC 2: Parameter & Math Function Selection ──────────────────────────────
E2 = f"{BASE}/Epic-2-Parameter-and-Math-Function-Selection"

create_us_doc(
    f"{E2}/US-03-Display-Parameter-Tab-with-Fields.docx",
    "US-03", "Display Parameter Tab with Available Fields",
    "Epic 2 - Parameter and Math Function Selection", "Frontend",
    "user",
    "click on the 'Parameter' tab in the expression builder popup",
    "I can view all available fields to use in my expression",
    [
        "Clicking 'Parameter' tab displays only the configured fields: FieldA, FieldB, FieldC, FieldD, FieldE, FieldF.",
        "Each field is displayed as a selectable option.",
        "An additional option called 'Number' is also displayed under Parameter.",
        "No other options are shown when 'Parameter' tab is active.",
        "The list is scrollable if the number of fields exceeds visible area.",
    ]
)

create_us_doc(
    f"{E2}/US-04-Filter-Search-Parameters.docx",
    "US-04", "Filter/Search Parameters by Field Name",
    "Epic 2 - Parameter and Math Function Selection", "Frontend",
    "user",
    "type a field name in the search box within the Parameter tab",
    "I can quickly find the field I need without scrolling through all options",
    [
        "A search/filter text box is present within the Parameter tab.",
        "As the user types, the field list is filtered in real-time to show only matching fields.",
        "The search is case-insensitive.",
        "If no fields match the search input, an appropriate empty/no-results message is displayed.",
        "Clearing the search box restores the full list of fields.",
        "The 'Number' option is also searchable/filterable.",
    ],
    ["Needs Clarification: Should search filter the 'Number' option as well, or only named fields?"]
)

create_us_doc(
    f"{E2}/US-05-Display-Math-Function-Options.docx",
    "US-05", "Display Math Function Options",
    "Epic 2 - Parameter and Math Function Selection", "Frontend",
    "user",
    "click on the 'Math Function' tab in the expression builder popup",
    "I can see all available math functions to apply in my expression",
    [
        "Clicking 'Math Function' tab displays exactly the following options: Min, MAX, FLOOR, ABS, BRACKET, CEIL.",
        "No other options are shown when 'Math Function' tab is active.",
        "Each math function is displayed as a selectable option.",
        "The label 'Math Function' remains unchanged.",
    ]
)

create_us_doc(
    f"{E2}/US-06-Display-All-Options.docx",
    "US-06", "Display All Options via All Tab",
    "Epic 2 - Parameter and Math Function Selection", "Frontend",
    "user",
    "click on the 'All' tab in the expression builder popup",
    "I can see all available parameters and math functions in one combined view",
    [
        "Clicking 'All' tab displays both all fields (FieldA-FieldF, Number) and all math functions (Min, MAX, FLOOR, ABS, BRACKET, CEIL).",
        "The combined list is displayed in a clear, readable format.",
        "The 'All' label remains unchanged.",
        "Switching from 'Parameter' or 'Math Function' tab to 'All' refreshes the list immediately.",
    ]
)

create_us_doc(
    f"{E2}/US-07-Add-Parameter-Field-to-Expression.docx",
    "US-07", "Add Parameter Field to Add Expression Area",
    "Epic 2 - Parameter and Math Function Selection", "Frontend",
    "user",
    "click on a field (e.g., FieldA) from the Parameter list",
    "that field is added as a token in the 'Add Expression' area along with a '+' sign",
    [
        "Clicking on a field from the Parameter list adds that field as a token in the 'Add Expression' area.",
        "A '+' (plus) sign/button is added immediately after the field token in the expression area.",
        "The field token is visually distinct (e.g., styled chip/badge).",
        "Multiple fields can be added to the expression.",
    ],
    ["Needs Clarification: Does the popup close after selecting a field, or stay open for further selection?"]
)

# ─── EPIC 3: Number Input in Expression ─────────────────────────────────────
E3 = f"{BASE}/Epic-3-Number-Input-in-Expression"

create_us_doc(
    f"{E3}/US-08-Add-Number-to-Expression.docx",
    "US-08", "Add Number Placeholder to Expression",
    "Epic 3 - Number Input in Expression", "Frontend",
    "user",
    "click on 'Number' from the Parameter list",
    "a default numeric placeholder '0.0' is added to the expression area",
    [
        "Clicking 'Number' in the Parameter tab adds '0.0' as a token in the 'Add Expression' area.",
        "The '0.0' token is visually distinguishable from field tokens.",
        "A '+' sign is added after '0.0' just like other tokens.",
        "The '0.0' token is clickable for editing.",
    ]
)

create_us_doc(
    f"{E3}/US-09-Edit-Number-via-Popup.docx",
    "US-09", "Edit Numeric Value via Popup",
    "Epic 3 - Number Input in Expression", "Frontend",
    "user",
    "click on the '0.0' token in the expression area",
    "I can enter a custom numeric value to replace the placeholder",
    [
        "Clicking '0.0' in the expression area opens a popup for numeric input.",
        "The popup contains an input field pre-filled with the current value ('0.0').",
        "The popup contains a 'Save' button.",
        "Clicking 'Save' replaces '0.0' with the entered value (e.g., '10.0') in the expression area.",
        "The updated value is reflected immediately in the expression.",
        "Input validation: only valid numeric values should be accepted.",
        "Clicking outside the popup or a Cancel button dismisses it without saving.",
    ],
    [
        "Needs Clarification: Should decimal-only numbers be allowed, or integers as well?",
        "Needs Clarification: Is there a Cancel button in the number edit popup?",
    ]
)

# ─── EPIC 4: Operator & Bracket Expression Building ─────────────────────────
E4 = f"{BASE}/Epic-4-Operator-and-Bracket-Expression-Building"

create_us_doc(
    f"{E4}/US-10-Add-Math-Operator-Between-Operands.docx",
    "US-10", "Add Math Operator Between Operands via Plus Sign",
    "Epic 4 - Operator and Bracket Expression Building", "Frontend",
    "user",
    "click on the '+' (plus) sign after a token in the expression area",
    "I can select a math operator to connect two operands in my expression",
    [
        "Clicking the '+' sign after a token opens an operator selection popup.",
        "The popup displays math operator options.",
        "Selecting an operator adds it to the expression between the current token and a new operand slot.",
        "After selecting an operator, the Parameter/Math Function/All dropdown appears for the next operand.",
        "The flow continues: user selects operand, '+' sign appears, user can add more operators.",
    ],
    ["Needs Clarification: What are the exact math operators available in the operator popup (+, -, *, /)?"]
)

create_us_doc(
    f"{E4}/US-11-Add-Bracket-Expression.docx",
    "US-11", "Add Bracket Grouped Sub-Expression",
    "Epic 4 - Operator and Bracket Expression Building", "Frontend",
    "user",
    "select 'BRACKET' from the Math Function options",
    "the expression area displays a bracket structure for grouped sub-expressions",
    [
        "Selecting 'BRACKET' from Math Function adds a bracket structure to the 'Add Expression' area.",
        "The bracket is visually represented with opening '(' and closing ')' tokens.",
        "Inside the bracket, a '+' sign is available for adding operands.",
        "Outside the bracket, a '+' sign is available for adding further operands.",
        "The bracket structure is visually distinct from regular tokens.",
    ]
)

create_us_doc(
    f"{E4}/US-12-Add-Left-Operand-Inside-Bracket.docx",
    "US-12", "Add Left Operand Inside Bracket via Plus Sign",
    "Epic 4 - Operator and Bracket Expression Building", "Frontend",
    "user",
    "click on the '+' sign inside a bracket",
    "I can select a parameter or field to act as the left operand inside the bracket",
    [
        "Clicking the '+' inside the bracket opens the Parameter/Math Function/All dropdown popup.",
        "Selecting a field adds it as the left operand inside the bracket.",
        "After selecting the left operand, a '+' sign appears for adding the operator.",
        "The selected field is displayed inside the bracket in the expression area.",
    ]
)

create_us_doc(
    f"{E4}/US-13-Add-Operator-and-Right-Operand-Inside-Bracket.docx",
    "US-13", "Add Operator and Right Operand Inside Bracket",
    "Epic 4 - Operator and Bracket Expression Building", "Frontend",
    "user",
    "click the '+' sign after the left operand inside a bracket",
    "I can add an operator and then a right operand to complete the bracketed sub-expression",
    [
        "Clicking '+' after the left operand inside the bracket opens the operator popup.",
        "Selecting an operator adds it inside the bracket.",
        "After the operator, the Parameter/Math Function/All dropdown appears for the right operand.",
        "Selecting the right operand completes the bracketed sub-expression.",
        "Additional operators/operands can be added inside the bracket by clicking '+' again.",
        "The completed bracket expression is displayed as: (LeftOperand Operator RightOperand).",
    ]
)

create_us_doc(
    f"{E4}/US-14-Add-Operand-Operator-Outside-Bracket.docx",
    "US-14", "Add Operator and Operand Outside Bracket",
    "Epic 4 - Operator and Bracket Expression Building", "Frontend",
    "user",
    "click the '+' sign outside the bracket in the expression area",
    "I can extend the expression beyond the bracket with additional operators and operands",
    [
        "A '+' sign is visible outside the bracket in the expression area.",
        "Clicking the '+' outside the bracket opens the operator popup first.",
        "After selecting the operator, the Parameter/Math Function/All dropdown appears for the next operand.",
        "Selecting the operand appends it to the expression outside the bracket.",
        "The flow continues: each '+' click allows adding more operator-operand pairs.",
    ]
)

# ─── EPIC 5: Comparator Configuration ───────────────────────────────────────
E5 = f"{BASE}/Epic-5-Comparator-Configuration"

create_us_doc(
    f"{E5}/US-15-Add-Comparator-to-Condition.docx",
    "US-15", "Add Comparator to Condition",
    "Epic 5 - Comparator Configuration", "Frontend",
    "user",
    "click on the 'Add Comparator' label on the main screen",
    "I can select a comparison operator to define the condition comparison logic",
    [
        "Clicking 'Add Comparator' opens a popup with the following options: Equal to, Lower than, Lower than equal to, Higher than, Higher than equal to.",
        "Exactly these 5 options are displayed, no more, no less.",
        "Each option is clearly labeled with its comparison meaning.",
        "The popup is displayed in a clear, accessible format.",
    ]
)

create_us_doc(
    f"{E5}/US-16-Select-Comparator-Option.docx",
    "US-16", "Select Comparator Option and Add to Condition",
    "Epic 5 - Comparator Configuration", "Frontend",
    "user",
    "click on one of the comparator options",
    "the selected comparator is added to the condition expression area",
    [
        "Clicking one of the comparator options (e.g., 'Equal to') closes the popup.",
        "The selected comparator token is added/displayed in the 'Add Comparator' area on the main screen.",
        "The comparator token is visually distinct (e.g., highlighted chip/badge).",
        "Only one comparator can be selected per condition.",
        "The selected comparator can be changed by clicking 'Add Comparator' again.",
    ],
    ["Needs Clarification: Can the user change the comparator after it has been set, or must they delete the condition?"]
)

create_us_doc(
    f"{E5}/US-17-Add-Right-Side-Expression-to-Comparator.docx",
    "US-17", "Add Right-Side Expression After Comparator",
    "Epic 5 - Comparator Configuration", "Frontend",
    "user",
    "have a comparator selected in my condition",
    "I can build the right-hand side expression using the same expression builder",
    [
        "After a comparator is selected, an 'Add Expression' area/button is available on the right side of the comparator.",
        "Clicking this 'Add Expression' opens the same expression builder popup (Parameter/Math Function/All).",
        "The right-side expression follows the same building rules as the left-side expression.",
        "The full condition is displayed as: [Left Expression] [Comparator] [Right Expression].",
        "Both left and right expressions support fields, numbers, brackets, and math functions.",
    ]
)

# ─── EPIC 6: Condition Management ───────────────────────────────────────────
E6 = f"{BASE}/Epic-6-Condition-Management-AND-OR"

create_us_doc(
    f"{E6}/US-18-Delete-a-Condition.docx",
    "US-18", "Delete a Condition",
    "Epic 6 - Condition Management (AND/OR)", "Frontend",
    "user",
    "click the delete (bin) icon on a condition",
    "the condition is removed and the screen returns to its initial state",
    [
        "Each condition has a visible delete (bin/trash) icon.",
        "Clicking the delete icon removes the entire condition from the screen.",
        "If all conditions are deleted, the screen returns to the initial/empty state.",
        "The delete action is immediate.",
    ],
    ["Needs Clarification: Should a confirmation prompt appear before deleting a condition?"]
)

create_us_doc(
    f"{E6}/US-19-Add-New-AND-Condition.docx",
    "US-19", "Add New AND Condition via and/or Button",
    "Epic 6 - Condition Management (AND/OR)", "Frontend",
    "user",
    "click the '+ and/or' button",
    "a new condition line is added below the current condition with an 'AND' operator",
    [
        "Clicking '+ and/or' adds a new empty condition row below the existing condition.",
        "The new condition row is connected to the previous with an 'AND' operator by default.",
        "The new condition row has the same structure: 'Add Expression', 'Add Comparator', 'Add Expression', and a delete icon.",
        "Multiple condition rows can be added.",
        "Each new condition is appended below the previous one.",
    ]
)

create_us_doc(
    f"{E6}/US-20-Toggle-AND-OR-Operator.docx",
    "US-20", "Toggle AND/OR Operator Between Conditions",
    "Epic 6 - Condition Management (AND/OR)", "Frontend",
    "user",
    "click on the 'AND' label between two conditions",
    "it toggles to 'OR', and clicking 'OR' toggles it back to 'AND'",
    [
        "The operator label between conditions is a toggle button displaying either 'AND' or 'OR'.",
        "Default state is 'AND' when a new condition is added.",
        "Clicking 'AND' changes it to 'OR'.",
        "Clicking 'OR' changes it back to 'AND'.",
        "The toggle is visually distinct (e.g., button with active state styling).",
        "The toggle change is reflected immediately without page reload.",
    ]
)

# ─── EPIC 7: Group & Subgroup Conditions ────────────────────────────────────
E7 = f"{BASE}/Epic-7-Group-and-Subgroup-Conditions"

create_us_doc(
    f"{E7}/US-21-Create-Condition-Subgroup.docx",
    "US-21", "Create a Condition Subgroup",
    "Epic 7 - Group and Subgroup Conditions", "Frontend",
    "user",
    "click the 'Group' button",
    "a subgroup container is created to hold one or more conditions as a logical unit",
    [
        "Clicking 'Group' creates a visually enclosed subgroup container.",
        "The subgroup contains its own 'Add Expression', 'Add Comparator', and '+ and/or' controls.",
        "Conditions inside a group are logically grouped with parentheses in the rule evaluation.",
        "A delete icon is available to remove the entire group.",
        "The group is visually distinct from top-level conditions (e.g., indented or bordered).",
    ]
)

create_us_doc(
    f"{E7}/US-22-Nested-Subgroup-Support.docx",
    "US-22", "Support Nested Subgroups at Any Level",
    "Epic 7 - Group and Subgroup Conditions", "Frontend",
    "user",
    "click 'Group' within an existing subgroup",
    "a nested subgroup is created inside the parent group, supporting any level of nesting",
    [
        "Clicking 'Group' inside an existing subgroup creates a nested subgroup.",
        "Nesting can be done to any depth (unlimited levels).",
        "Each nested group is visually indented/differentiated from its parent.",
        "Each nested group has its own controls: 'Add Expression', 'Add Comparator', '+ and/or', 'Group', and delete.",
        "The logical structure (parenthesized nesting) is preserved at every level.",
    ],
    ["Needs Clarification: Is there a maximum nesting depth limit for performance/UI reasons?"]
)

# ─── EPIC 8: Backend API & Data Model ───────────────────────────────────────
E8 = f"{BASE}/Epic-8-Backend-API-and-Data-Model"

create_us_doc(
    f"{E8}/US-23-Expression-Data-Model-and-Persistence-API.docx",
    "US-23", "Expression Data Model and Persistence API",
    "Epic 8 - Backend API and Data Model", "Backend",
    "system",
    "persist a rule expression to the database via a REST API",
    "the configured expressions and conditions can be stored and retrieved for future use",
    [
        "A REST API endpoint exists to save an RBP rule expression (e.g., POST /api/rbp-rules).",
        "The data model supports: fields/parameters, math functions, operators, comparators, AND/OR logic, and nested groups.",
        "The API accepts a structured JSON representation of the full rule expression.",
        "The API returns the saved rule with a unique identifier.",
        "A GET endpoint exists to retrieve saved rules by ID.",
        "The data model supports unlimited nesting of subgroups.",
        "Input validation is enforced at the API level.",
    ],
    [
        "Needs Clarification: Is there a requirement to save/load rules, or is this a session-only UI tool?",
        "Needs Clarification: What is the API base path and authentication mechanism?",
    ]
)

create_us_doc(
    f"{E8}/US-24-Validate-Expression-Structure.docx",
    "US-24", "Validate Expression Structure on the Backend",
    "Epic 8 - Backend API and Data Model", "Backend",
    "system",
    "validate the structure of an RBP expression submitted via the API",
    "only syntactically and semantically valid rule expressions are stored",
    [
        "The backend validates that each condition has a left expression, a comparator, and a right expression.",
        "Expressions must contain at least one valid operand (field, number, or math function result).",
        "Nested bracket expressions must be properly formed (left and right operands with operator).",
        "AND/OR operators must connect exactly two conditions or groups.",
        "Invalid expressions return a 400 Bad Request with descriptive error messages.",
        "Field names in expressions are validated against the configured fields (FieldA-FieldF, Number).",
    ],
    [
        "Needs Clarification: Are field names (FieldA-FieldF) configurable per-project or hardcoded?",
        "Needs Clarification: Should frontend also perform inline expression validation before submission?",
    ]
)

print("\nAll 24 user story DOCX files created successfully!")
